"""
Word document formatter for OCR results.
Converts structured ContentBlocks to Word document format.
"""

from typing import List, Dict, Optional
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image
import os
import re
from html.parser import HTMLParser


class HTMLTableParser(HTMLParser):
    """Parse HTML table to extract rows and cells."""

    def __init__(self):
        super().__init__()
        self.tables = []
        self.current_table = []
        self.current_row = []
        self.current_cell = ""
        self.in_table = False
        self.in_row = False
        self.in_cell = False

    def handle_starttag(self, tag, attrs):
        if tag == 'table':
            self.in_table = True
            self.current_table = []
        elif tag == 'tr':
            self.in_row = True
            self.current_row = []
        elif tag in ['td', 'th']:
            self.in_cell = True
            self.current_cell = ""

    def handle_endtag(self, tag):
        if tag == 'table':
            self.in_table = False
            if self.current_table:
                self.tables.append(self.current_table)
                self.current_table = []
        elif tag == 'tr':
            self.in_row = False
            if self.current_row:
                self.current_table.append(self.current_row)
                self.current_row = []
        elif tag in ['td', 'th']:
            self.in_cell = False
            self.current_row.append(self.current_cell.strip())
            self.current_cell = ""

    def handle_data(self, data):
        if self.in_cell:
            self.current_cell += data


class WordFormatter:
    """Format OCR results as Word document with proper styling."""

    def __init__(self, output_dir: str):
        """
        Initialize Word formatter.

        Args:
            output_dir: Directory to save Word documents and images
        """
        self.output_dir = output_dir
        self.images_dir = os.path.join(output_dir, "images")
        os.makedirs(self.images_dir, exist_ok=True)

        self.doc = Document()
        self.image_counter = 0
        self.current_image = None

        # 设置页面格式和默认样式
        self._setup_page_format()
        self._setup_default_styles()
        self._setup_page_numbers()

    def _setup_page_format(self):
        """设置页面格式（A4纸张，标准页边距）"""
        # 获取文档的sections
        for section in self.doc.sections:
            # 设置A4纸张大小
            section.page_height = Cm(29.7)  # A4高度
            section.page_width = Cm(21.0)   # A4宽度

            # 设置页边距（上下2.54cm，左右3.17cm）
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(3.17)
            section.right_margin = Cm(3.17)

    def _setup_default_styles(self):
        """设置默认字体样式"""
        # 设置默认正文样式
        style = self.doc.styles['Normal']
        font = style.font
        font.name = '宋体'  # 中文字体
        font.size = Pt(10.5)  # 标准正文字号
        # 设置中文字体
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # 设置行距为1.5倍
        paragraph_format = style.paragraph_format
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    def _setup_page_numbers(self):
        """在页脚添加页码"""
        section = self.doc.sections[0]
        footer = section.footer

        # 在页脚添加段落
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 添加页码域
        run = paragraph.add_run()

        # 创建页码域代码
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')

        run._element.append(fldChar1)
        run._element.append(instrText)
        run._element.append(fldChar2)

        # 设置页码字体格式
        run.font.name = '宋体'
        run.font.size = Pt(10.5)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    def _clean_escape_chars(self, text: str) -> str:
        """
        清理文本中的LaTeX/Markdown转义字符

        Args:
            text: 原始文本

        Returns:
            清理后的文本
        """
        if not text:
            return text

        # 清理LaTeX/Markdown数学公式标记
        text = re.sub(r'\\\(', '(', text)  # \( -> (
        text = re.sub(r'\\\)', ')', text)  # \) -> )
        text = re.sub(r'\\\[', '[', text)  # \[ -> [
        text = re.sub(r'\\\]', ']', text)  # \] -> ]
        text = re.sub(r'\\\{', '{', text)  # \{ -> {
        text = re.sub(r'\\\}', '}', text)  # \} -> }
        text = re.sub(r'\\%', '%', text)   # \% -> %
        text = re.sub(r'\\_', '_', text)   # \_ -> _
        text = re.sub(r'\\#', '#', text)   # \# -> #
        text = re.sub(r'\\&', '&', text)   # \& -> &
        text = re.sub(r'\\\$', '$', text)  # \$ -> $

        return text

    def format_blocks(self, blocks: List[Dict], source_image: Optional[Image.Image] = None) -> Document:
        """
        Format content blocks into Word document.

        Args:
            blocks: List of ContentBlock dictionaries
            source_image: Source image for extracting image blocks

        Returns:
            Document object
        """
        self.current_image = source_image

        for block in blocks:
            block_type = block.get("type", "text")
            # Convert to uppercase for consistency (MinerU returns lowercase)
            block_type = block_type.upper() if block_type else "TEXT"

            content = block.get("content")

            # Handle None content
            if content is None:
                content = ""
            else:
                content = content.strip()
                # Clean escape characters for Word output
                content = self._clean_escape_chars(content)

            if not content and block_type not in ["IMAGE"]:
                continue

            if block_type == "TITLE":
                self._add_title(content)
            elif block_type == "TEXT":
                self._add_paragraph(content)
            elif block_type == "IMAGE":
                self._add_image(block)
            elif block_type == "TABLE":
                self._add_table(content)
            elif block_type == "CODE":
                self._add_code(content)
            elif block_type == "EQUATION":
                self._add_equation(content)
            elif block_type == "PAGE_NUMBER":
                # 跳过PAGE_NUMBER，使用页脚的页码替代
                continue
            elif block_type in ["HEADER", "FOOTER"]:
                self._add_header_footer(content, block_type)
            else:
                # Default: treat as paragraph
                self._add_paragraph(content)

        return self.doc

    def _add_title(self, content: str):
        """Add a title/heading to the document."""
        heading = self.doc.add_heading(content, level=2)
        # 设置标题字体和大小
        for run in heading.runs:
            run.font.name = '黑体'  # 标题使用黑体
            run.font.size = Pt(14)  # 二级标题14pt
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
            # 设置中文字体
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

        # 设置段落格式
        heading.paragraph_format.space_before = Pt(12)  # 段前间距
        heading.paragraph_format.space_after = Pt(6)    # 段后间距

    def _add_paragraph(self, content: str):
        """Add a normal paragraph to the document."""
        if not content:
            return
        paragraph = self.doc.add_paragraph(content)
        # Use default style (宋体 10.5pt) set in _setup_default_styles()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    def _add_image(self, block: Dict):
        """Add an image to the document."""
        if not self.current_image:
            print(f"Warning: No current_image available for IMAGE block")
            return

        # Save image from bbox
        image_path = self._save_image_from_bbox(block)
        if not image_path:
            print(f"Warning: Could not save image from bbox")
            return

        if not os.path.exists(image_path):
            print(f"Warning: Image file does not exist: {image_path}")
            return

        try:
            # Add image to document
            paragraph = self.doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Calculate appropriate width (max 6 inches)
            img = Image.open(image_path)
            aspect_ratio = img.height / img.width
            width = min(6.0, img.width / 100)  # Reasonable width

            run = paragraph.add_run()
            run.add_picture(image_path, width=Inches(width))

            # Add caption if available
            caption = block.get("content", "")
            if caption:
                caption_para = self.doc.add_paragraph(caption)
                caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Apply caption formatting
                for run in caption_para.runs:
                    run.font.name = '宋体'
                    run.font.size = Pt(9)
                    run.font.italic = True
                    run.font.color.rgb = RGBColor(64, 64, 64)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

            print(f"Successfully added image: {image_path}")
        except Exception as e:
            print(f"Warning: Could not insert image {image_path}: {e}")
            import traceback
            traceback.print_exc()

    def _add_table(self, content: str):
        """Add a table to the document."""
        rows = []

        # Check if it's HTML table
        if '<table>' in content.lower():
            # Parse HTML table
            parser = HTMLTableParser()
            parser.feed(content)

            if parser.tables and len(parser.tables) > 0:
                rows = parser.tables[0]  # Get first table
        else:
            # Parse markdown table
            lines = content.strip().split('\n')
            if len(lines) < 2:
                # If can't parse as table, add as paragraph
                self._add_paragraph(content)
                return

            # Remove separator line (usually second line with |---|---|)
            table_lines = [line for line in lines if not all(c in '|-: ' for c in line)]

            if not table_lines:
                return

            # Parse table structure
            for line in table_lines:
                cells = [cell.strip() for cell in line.split('|')]
                # Remove empty first/last cells from | ... |
                cells = [c for c in cells if c]
                if cells:
                    rows.append(cells)

        if not rows:
            return

        # Create Word table
        num_cols = len(rows[0])
        table = self.doc.add_table(rows=len(rows), cols=num_cols)
        table.style = 'Light Grid Accent 1'

        # Fill table
        for i, row in enumerate(rows):
            for j, cell_text in enumerate(row):
                if j < len(table.rows[i].cells):
                    cell = table.rows[i].cells[j]
                    # Clean escape characters in cell text
                    cell.text = self._clean_escape_chars(cell_text)
                    # Apply font formatting to cell content
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = '宋体'
                            run.font.size = Pt(10.5)
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                            # Header row formatting
                            if i == 0:
                                run.font.bold = True

        # Add spacing after table
        self.doc.add_paragraph()

    def _add_code(self, content: str):
        """Add code block to the document."""
        paragraph = self.doc.add_paragraph()
        paragraph.style = 'No Spacing'

        # Set monospace font
        run = paragraph.add_run(content)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)

        # Set font for Asian text (for better compatibility)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')

        # Light gray background (simulate code block)
        shading_elm = run._element.get_or_add_rPr().get_or_add_shd()
        shading_elm.set(qn('w:fill'), 'F0F0F0')

    def _add_equation(self, content: str):
        """Add mathematical equation to the document."""
        # Add equation as text with special formatting
        # For better support, could use Office Math XML in the future
        paragraph = self.doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = paragraph.add_run(content)
        run.font.italic = True
        run.font.size = Pt(11)

    def _add_header_footer(self, content: str, block_type: str):
        """Add header/footer comment to the document."""
        paragraph = self.doc.add_paragraph(f"[{block_type}]: {content}")
        # Apply header/footer formatting
        for run in paragraph.runs:
            run.font.name = '宋体'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(128, 128, 128)
            run.font.italic = True
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    def _save_image_from_bbox(self, block: Dict) -> Optional[str]:
        """Extract and save image from bbox coordinates."""
        if not self.current_image:
            return None

        bbox = block.get("bbox")
        if not bbox or len(bbox) != 4:
            return None

        self.image_counter += 1
        image_filename = f"image_{self.image_counter}.jpg"
        image_path = os.path.join(self.images_dir, image_filename)

        try:
            # Get image dimensions
            img_width, img_height = self.current_image.size

            # Convert normalized bbox to pixel coordinates
            x1 = int(bbox[0] * img_width)
            y1 = int(bbox[1] * img_height)
            x2 = int(bbox[2] * img_width)
            y2 = int(bbox[3] * img_height)

            # Ensure coordinates are valid
            x1, x2 = max(0, min(x1, x2)), min(img_width, max(x1, x2))
            y1, y2 = max(0, min(y1, y2)), min(img_height, max(y1, y2))

            if x2 <= x1 or y2 <= y1:
                return None

            # Crop image
            cropped = self.current_image.crop((x1, y1, x2, y2))

            # Convert to RGB if necessary (JPEG doesn't support RGBA)
            if cropped.mode in ('RGBA', 'LA', 'P'):
                rgb_image = Image.new('RGB', cropped.size, (255, 255, 255))
                if cropped.mode == 'P':
                    cropped = cropped.convert('RGBA')
                rgb_image.paste(cropped, mask=cropped.split()[-1] if cropped.mode in ('RGBA', 'LA') else None)
                cropped = rgb_image

            # Save image
            cropped.save(image_path, "JPEG", quality=95)
            return image_path

        except Exception as e:
            print(f"Warning: Could not save image from bbox: {e}")
            return None

    def save(self, output_path: str):
        """Save the Word document to file."""
        self.doc.save(output_path)
